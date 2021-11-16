import docx

'''d = docx.Document('test.docx')
print(d.paragraphs[1].text)
p = d.paragraphs[1]

for run in p.runs:
    print(run.text)'''

d = docx.Document()
d.add_paragraph('Hello this is a paragraph.')
d.add_paragraph('Hello this is another paragraph.')

p = d.paragraphs[0]
p.add_run('this is a new run')
p.runs[1].bold = True
d.save('test4.docx')
